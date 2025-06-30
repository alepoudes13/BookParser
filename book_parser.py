try:
    from search import *

    from bs4 import BeautifulSoup
    import urllib.request, requests
    from docx import Document #pip install python-docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor
    from docx.text.run import Font
    import os, sys, math, time
    if not(sys.platform == "linux" or sys.platform == "linux2" or sys.platform == "darwin"):
        import comtypes.client
    with open(os.devnull, 'w') as f:
        oldstdout = sys.stdout
        sys.stdout = f

        import pygame

        sys.stdout = oldstdout


    b = False
    it = False
    div = False
    p = False

    output = None
    para = None

    def dfs(s):
        global para, b, it, div, p, output
        atr = ''
        if s.name == 'a':
            pass
        elif(s.name != None):
            atr = ' '.join(s.attrs.get('class', [])).split()
            if s.name == u'img':
                r = para.add_run('\n')
                path = 'img/' + s['src'].split('/')[-1]
                img = pygame.image.load(path)
                width = img.get_width()
                height = img.get_height()
                if width > 500:
                    dcc = height / width
                    width = 500
                    height = math.floor(dcc * width)
                if height > 700:
                    dcc = width / height
                    height = 700
                    width = math.floor(dcc * height)
                r.add_picture(path, width = Pt(width), height = Pt(height))
            if s.name == u'br':
                para.add_run('\n')
            if s.name == u'p':
                p = True
                para = output.add_paragraph()
            if s.name == u'div' and ('take_h1' in atr):
                div = True
                b = True
                para = output.add_paragraph()
            if 'strong' in atr:
                b = True
            if 'em' in atr:
                it = True
            if s.name == u'b':
                b = True
            if s.name == u'i':
                it = True
            for i in s.contents:
                dfs(i)
            if s.name == u'p':
                p = False
            if s.name == u'div' and ('take_h1' in atr):
                div = False
                b = False
            if 'strong' in atr:
                b = False
            if 'em' in atr:
                it = False
            if s.name == u'b':
                b = False
            if s.name == u'i':
                it = False
        else:
            run = None
            if div == True and ('take_h1' in ' '.join(s.parent.attrs.get('class', [])).split()):
                run = para.add_run(' '.join(str(s.string).split()))
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.font.color.rgb = RGBColor(0x00, 0x3a, 0xac)
            else:
                t = str(s.string)
                rt = ''
                for i in t:
                    if i != '\n' and i != '\r':
                        rt += i
                    if i == '\r':
                        rt += ' '
                run = para.add_run(rt)
            if b == True:
                run.bold = True
            if it == True:
                run.italic = True

    #============

    print("""Hello! It's book_parser 3.0. Enjoy the reading!
<---------------------------------->""")

    mode = input("Choose mode: 1 - searching, 2 - by id's: ")

    try:
        os.makedirs("img")
    except:
        pass
    try:
        os.makedirs("DOC")
    except:
        pass
    try:
        os.makedirs("PDF")
    except:
        pass

    nm = []
    idm = []

    if mode == "2":

        print("""====================================
Example:
id: 001 002
save as: name1; name2
====================================""")
        idm = input('id: ').split()
        iid = 0
        while iid < len(idm):
            lst = input('save as: ').split(';')
            for i in lst:
                nm.append(i)
            iid += len(lst)

    else:
        print("====================================")
        [idm, nm] = search()
        print("====================================")

    dec = input("Convert to PDF automatically(Y,Д/N,Н)? Works only on Windows: ")
    acc = ['Y', 'y', 'Д', 'д']
    convert = True if dec in acc else False



    for id in range(len(idm)):
        output = Document()
        para = output.add_paragraph().clear()
    ##try:
    ##    a, e = map(int, input('from page _ to _: ').split())
    ##except:
        html_doc = urllib.request.urlopen('http://loveread.ec/read_book.php?id=' + idm[id] + '&p=1').read()
        soup = BeautifulSoup(html_doc, 'html.parser')

        s = soup.find('div', 'navigation')
        a, e = (1, int(s.findAll('a')[-2]['href'].split('=')[-1]))

        for i in range(a, e + 1):
            html_doc = urllib.request.urlopen('http://loveread.ec/read_book.php?id=' + idm[id] + '&p=' + str(i)).read()
            soup = BeautifulSoup(html_doc, 'html.parser')

            s = soup.find('div', 'MsoNormal')

            images = s.findAll('img')
            for image in images:
                adr = image['src']
                pict = requests.get('http://loveread.ec/' + adr)
                out = open('img/' + adr.split('/')[-1], "wb")
                out.write(pict.content)
                out.close()

            for k in s.contents[2:-2]:
                dfs(k)

            print('[System]| page ' + str(i - a + 1) + ' of ' + str(e - a + 1) + ' is copied (' + str(id + 1) + ' of ' + str(len(idm)) + ') |[System]')

        print("Saving as docx...")
        style = output.styles['Normal']
        font = style.font
        font.size = Pt(22)

        forb = ['/', '\\', '?', '*', '"', '<', '>', '|', ':']
        word_name = ''
        for i in nm[id]:
            if i in forb:
                continue
            word_name += i

        try:
            output.save("DOC/" + word_name + '.docx')
        except:
            word_name += '(1)'
            output.save("DOC/" + word_name + '.docx')

        if convert == True:
            if not(sys.platform == "linux" or sys.platform == "linux2" or sys.platform == "darwin"):
                print("Saving as pdf...")
                pth = os.path.abspath("DOC\\" + word_name + '.docx').split("DOC\\" + word_name + '.docx')[0]
                PDF = 17
                word = comtypes.client.CreateObject('Word.Application')
                time.sleep(3)

                doc = word.Documents.Open(pth + "DOC/" + word_name + '.docx')
                doc.SaveAs(pth + '/pdf/' + word_name + '.pdf', FileFormat = PDF)
                doc.Close()
                word.Quit()
            else:
                print("/----PDF works only on Windows")

        print('====================================================')

    input("Press ENTER to continue...")
except Exception as e:
    print(e)
    input("Error...")
